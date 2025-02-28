import streamlit as st
import pandas as pd
import openai
import os
import tempfile
import zipfile
from docx import Document

st.set_page_config(page_title="Content Generation")

st.title("📄 Generate Content from Excel")

# 📌 Inserimento della chiave API di OpenAI
oai_api_key = st.text_input("Insert your OpenAI API KEY", type="password")
if not oai_api_key:
    st.warning("⚠️ Insert your API KEY to continue.")
    st.stop()

oai_client = openai.Client(api_key=oai_api_key)

# 📌 Caricamento del file Excel
uploaded_file = st.file_uploader("Upload Excel file", type=["xlsx"])

if uploaded_file:
    df = pd.read_excel(uploaded_file)
    df.columns = df.columns.str.strip()
    st.write("📊 File Preview:")
    st.dataframe(df.head())

    temp_dir = tempfile.mkdtemp()

    # 📌 Generazione del contenuto
    def generate_paragraph(title, description, min_words):
        prompt = f"""
        Scrivi un paragrafo dettagliato con il seguente titolo e descrizione:
        **Titolo:** {title}
        **Descrizione:** {description}
        - Deve contenere almeno {min_words} parole.
        """

        response = oai_client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "system", "content": "Sei un copywriter esperto in SEO."},
                      {"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=3000
        )

        return response.choices[0].message.content

    if st.button("Generate Content"):
        generated_files = []

        for index, row in df.iterrows():
            st.write(f"📝 Generating content for: {row['Titolo']}...")
            min_words = row.get("Numero Minimo Parole", 800)
            full_text = "\n\n".join(
                f"{row[f'Paragrafo {i} (Titolo)']}\n\n{generate_paragraph(row[f'Paragrafo {i} (Titolo)'], row[f'Paragrafo {i} (Descrizione)'], min_words // 5)}"
                for i in range(1, 11)
                if f"Paragrafo {i} (Titolo)" in row and pd.notna(row[f"Paragrafo {i} (Titolo)"])
            )

            # 📌 Creazione del file Word
            file_path = os.path.join(temp_dir, f"{row['Titolo'].replace(' ', '_')}.docx")
            doc = Document()
            doc.add_heading(row['Titolo'], level=1)
            doc.add_paragraph(full_text)
            doc.save(file_path)
            generated_files.append(file_path)

        # 📌 Download dei file generati
        zip_path = os.path.join(temp_dir, "generated_content.zip")
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for file in generated_files:
                zipf.write(file, os.path.basename(file))

        with open(zip_path, "rb") as f:
            st.download_button("📥 Download Generated Content", f, file_name="generated_content.zip")