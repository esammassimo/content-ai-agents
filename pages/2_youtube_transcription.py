import streamlit as st
import openai
from youtube_transcript_api import YouTubeTranscriptApi
import re
import tempfile
import os
from docx import Document

st.set_page_config(page_title="YouTube Summarization")

st.title("🎬 Summarize YouTube Video")

oai_api_key = st.text_input("Insert your OpenAI API KEY", type="password")
if not oai_api_key:
    st.warning("⚠️ Insert your API KEY to continue.")
    st.stop()

oai_client = openai.Client(api_key=oai_api_key)

video_url = st.text_input("Enter YouTube video URL:")
min_words = st.number_input("Minimum number of words in summary:", min_value=50, value=100, step=10)
lang_code = st.text_input("Enter language code for transcript (default 'en'):", value="en")

def extract_video_id(url):
    match = re.search(r"(?:v=|youtu\.be/|embed/|/v/|/e/|watch\?v=|&v=)([a-zA-Z0-9_-]{11})", url)
    return match.group(1) if match else None

if st.button("Extract & Summarize"):
    video_id = extract_video_id(video_url)
    if not video_id:
        st.error("❌ Invalid YouTube URL. Please enter a valid link.")
        st.stop()

    transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=[lang_code])
    full_text = " ".join([entry['text'] for entry in transcript])
    st.text_area("Extracted Transcript:", full_text, height=300)

    def summarize_text(text):
        response = oai_client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "system", "content": "Sei un esperto nel sintetizzare contenuti di video."},
                      {"role": "user", "content": text}],
            temperature=0.5,
            max_tokens=3500
        )
        return response.choices[0].message.content

    summary = summarize_text(full_text)
    st.subheader("📌 Video Summary:")
    st.write(summary)

    temp_dir = tempfile.mkdtemp()
    summary_file = os.path.join(temp_dir, f"summary_{video_id}.docx")

    doc = Document()
    doc.add_heading("YouTube Video Summary", level=1)
    doc.add_paragraph(summary)
    doc.save(summary_file)

    with open(summary_file, "rb") as docx_file:
        st.download_button("📥 Download Summary (.docx)", docx_file, file_name=f"summary_{video_id}.docx")