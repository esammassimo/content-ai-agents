import streamlit as st
import openai
import os
import tempfile
from docx import Document

# 📌 Configurazione API OpenAI
st.set_page_config(page_title="AI Content Generator", layout="wide")
st.title("📝 AI Content Generator - Form Version")

# 📌 Inserimento della chiave API di OpenAI
oai_api_key = st.text_input("Insert your OpenAI API KEY", type="password")
if not oai_api_key:
    st.warning("⚠️ Insert your API KEY to continue.")
    st.stop()

oai_client = openai.Client(api_key=oai_api_key)

# 📌 Form per l'inserimento dei dati
with st.form("content_form"):
    st.header("📄 Content Details")
    title = st.text_input("Title of the content")
    seo_title = st.text_input("SEO Title")
    meta_description = st.text_area("Meta Description")
    min_words = st.number_input("Minimum number of words:", min_value=100, value=800, step=100)

    # 📌 Nuovo campo per il Tone of Voice
    tone_of_voice = st.selectbox("Tone of Voice", ["Informale", "Professionale", "Formale"])

    st.subheader("📝 Paragraphs (max 10)")
    paragraphs = []
    for i in range(1, 11):
        col1, col2 = st.columns([1, 2])
        with col1:
            p_title = st.text_input(f"Paragrafo {i} - Titolo", key=f"title_{i}")
        with col2:
            p_desc = st.text_area(f"Paragrafo {i} - Descrizione", key=f"desc_{i}")
        
        if p_title and p_desc:
            paragraphs.append((p_title, p_desc))

    submit_button = st.form_submit_button("Generate Content")

# 📌 Funzione per generare i paragrafi con OpenAI
def generate_paragraph(title, description, min_words, tone_of_voice):
    prompt = f"""
    Scrivi un paragrafo dettagliato con il seguente titolo e descrizione:
    **Titolo:** {title}
    **Descrizione:** {description}
    
    - Il tono di voce deve essere **{tone_of_voice.lower()}**.
    - Deve contenere almeno {min_words} parole.
    """

    response = oai_client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "Sei un copywriter esperto in SEO."},
                  {"role": "user", "content": prompt}],
        temperature=0.7,
        max_tokens=3000
    )

    return response.choices[0].message.content

# 📌 Se il form è stato inviato
if submit_button:
    if not title or not seo_title or not meta_description or len(paragraphs) == 0:
        st.error("❌ Please fill out all required fields.")
    else:
        st.write(f"📝 Generating content for: {title}...")

        # 📌 Generazione del contenuto
        full_text = []
        for p_title, p_desc in paragraphs:
            paragraph_text = generate_paragraph(p_title, p_desc, min_words // len(paragraphs), tone_of_voice)
            full_text.append(f"## {p_title}\n\n{paragraph_text}")

        # 📌 Creazione del file Word
        temp_dir = tempfile.mkdtemp()
        file_path = os.path.join(temp_dir, f"{title.replace(' ', '_')}.docx")

        doc = Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"SEO Title: {seo_title}")
        doc.add_paragraph(f"Meta Description: {meta_description}")
        doc.add_paragraph(f"Tone of Voice: {tone_of_voice}")

        for paragraph in full_text:
            doc.add_paragraph(paragraph)

        doc.save(file_path)

        # 📌 Download del file
        with open(file_path, "rb") as docx_file:
            st.download_button("📥 Download Content (.docx)", docx_file, file_name=f"{title.replace(' ', '_')}.docx")
